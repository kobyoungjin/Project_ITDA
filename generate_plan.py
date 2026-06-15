"""
ITDA 프로젝트 전체 개발 계획서 생성기 (2026-05-21 기준)
실행: python generate_plan.py
출력: ITDA_개발계획서_20260521.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "ITDA_개발계획서_20260528.docx"

# ── 색상 팔레트 ────────────────────────────────────────
C_TITLE   = (0x1F, 0x49, 0x7D)
C_H1      = (0x1F, 0x49, 0x7D)
C_H2      = (0x44, 0x72, 0xC4)
C_H3      = (0x26, 0x63, 0x85)
C_HDR_BG  = "DCE6F1"
C_DONE_BG = "E2EFDA"   # 완료 행 초록
C_TODO_BG = "FFF2CC"   # 예정 행 노랑
C_BLOCK   = "F4F4F4"


def _shd(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    colors = {1: C_H1, 2: C_H2, 3: C_H3}
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(*colors.get(level, C_H1))
    return h


def add_table(doc, headers, rows, col_widths=None, row_colors=None):
    """row_colors: list of hex strings per row, None = white"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shd(cell, C_HDR_BG)

    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = (row_colors[ri] if row_colors and row_colors[ri] else None)
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if bg:
                _shd(cell, bg)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), C_BLOCK)
    pPr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    pPr.append(ind)
    doc.add_paragraph()


def add_callout(doc, text, color="FFF9C4"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:right"), "180")
    pPr.append(ind)
    doc.add_paragraph()


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════
    # 표지
    # ══════════════════════════════════════════
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ITDA (잇다) 프로젝트")
    r.bold = True; r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(*C_TITLE)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("전체 개발 계획서")
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(*C_H2)

    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run(
        "농인↔청인 소통 보조 수어 번역 웹앱\n"
        "VIDEO-ONLY MODE 기반 최종 설계"
    )
    r.font.size = Pt(11)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = info.add_run(
        "팀명: ITDA (잇다) 팀\n"
        "팀원: 고병진  |  인기동  |  손영란  |  손현정"
    )
    ri.font.size = Pt(12)
    ri.bold = True
    ri.font.color.rgb = RGBColor(*C_H2)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 목차
    # ══════════════════════════════════════════
    add_heading(doc, "목차", 1)
    toc = [
        "1. 프로젝트 개요",
        "2. 기대효과",
        "3. 현재 개발 현황 (완료 항목)",
        "4. 시스템 아키텍처",
        "5. 기술 스택",
        "6. API 설계",
        "7. 데이터 전략",
        "8. 단계별 개발 계획",
        "9. 미결 이슈 및 리스크",
        "10. 역할 분담",
        "11. 체크리스트",
    ]
    for item in toc:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 1. 프로젝트 개요
    # ══════════════════════════════════════════
    add_heading(doc, "1. 프로젝트 개요", 1)
    add_table(doc,
        ["항목", "내용"],
        [
            ["프로젝트명", "ITDA (잇다) — 농인↔청인 소통 보조 수어 번역 웹앱"],
            ["팀명", "ITDA (잇다) 팀"],
            ["팀원", "고병진  |  인기동  |  손영란  |  손현정"],
            ["핵심 목표", "한국어 텍스트/음성 입력 → 수어 영상 자동 재생으로 의사소통 지원"],
            ["대상 사용자", "청인 (농인에게 수어로 말하고 싶은 사람)"],
            ["설계 방식", "VIDEO-ONLY MODE — Supabase Storage 영상 스트리밍"],
            ["직전 설계 변경", "3D 아바타(Three.js) 방식 폐기 → 실제 수어 영상 직접 재생으로 전환"],
        ],
        col_widths=[1.8, 4.2]
    )

    add_heading(doc, "1-1. 변경 배경", 2)
    doc.add_paragraph(
        "초기 설계는 Three.js 3D 아바타 + 원본 영상을 동시에 재생하는 DUAL MODE였습니다. "
        "그러나 아바타 렌더링 복잡도, 리타게팅 파이프라인 유지 비용, "
        "그리고 실제 수어 전달력의 차이를 고려하여 "
        "실제 수어사의 MP4 영상을 직접 재생하는 VIDEO-ONLY MODE로 전환하였습니다."
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 2. 기대효과
    # ══════════════════════════════════════════
    add_heading(doc, "2. 기대효과", 1)
    add_table(doc,
        ["구분", "기대효과", "세부 내용"],
        [
            ["사회적 효과",
             "소통 장벽 해소",
             "농인과 청인 간 언어 장벽을 낮추어 일상·공공 영역에서 원활한 의사소통 지원"],
            ["사회적 효과",
             "정보 접근성 향상",
             "병원·관공서·상점 등에서 수어 통역사 없이도 기본적인 의사소통 가능"],
            ["교육적 효과",
             "수어 학습 지원",
             "청인이 수어 영상을 직접 확인하며 수어를 자연스럽게 학습하는 환경 제공"],
            ["기술적 효과",
             "AI 수어 인식 고도화",
             "카메라 기반 MediaPipe·KNN 분류기로 실시간 수어 인식 기술 발전에 기여"],
            ["기술적 효과",
             "데이터 자산 구축",
             "1,111개 이상의 수어 영상 데이터베이스 및 RAG 검색 엔진 축적"],
            ["서비스 효과",
             "즉각적 응답성",
             "Supabase CDN 기반 영상 스트리밍으로 검색 후 1초 이내 영상 제공 목표"],
            ["서비스 효과",
             "확장 가능성",
             "80개 시나리오 단어에서 3,695개 수어 어휘로 확장 가능한 아키텍처 설계"],
        ],
        col_widths=[1.2, 1.8, 3.0]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 3. 현재 개발 현황
    # ══════════════════════════════════════════
    add_heading(doc, "3. 현재 개발 현황 (완료 항목)", 1)

    add_heading(doc, "3-1. 백엔드 완료", 2)
    add_table(doc,
        ["파일", "기능", "상태"],
        [
            ["api/main.py", "FastAPI 앱, CORS, 라우터 등록", "완료"],
            ["api/routers/sign_language.py", "POST /api/sign-language/search — RAG 수어 검색", "완료"],
            ["api/routers/video_proxy.py", "GET /api/video/url/{keyword} — 영상 URL 반환", "완료"],
            ["api/routers/ws_vision.py", "WS /api/ws/vision — 카메라 실시간 수어 인식", "완료"],
            ["api/routers/stt.py", "STT 어댑터 (음성→텍스트)", "완료"],
            ["api/routers/collect.py", "KNN 학습 데이터 수집", "완료"],
            ["api/services/rag_engine.py", "FAISS + sentence-transformers 검색 엔진", "완료"],
            ["api/services/supabase_video.py", "Supabase Storage REST API 연동 (aiohttp)", "완료"],
            ["api/services/knn_classifier.py", "카메라 수어 KNN 분류기", "완료"],
            ["api/core/config.py", "환경변수 설정 (SUPABASE_URL, SUPABASE_SERVICE_KEY)", "완료"],
            [".env.example", "환경변수 템플릿", "완료"],
        ],
        col_widths=[2.4, 2.8, 0.8],
        row_colors=[C_DONE_BG]*11
    )

    add_heading(doc, "3-2. 프론트엔드 완료", 2)
    add_table(doc,
        ["파일", "기능", "상태"],
        [
            ["frontend/index.html", "앱 레이아웃 — 카메라 영역 + 영상 영역 + 입력 영역", "완료"],
            ["frontend/js/translator.js", "텍스트/STT 입력 → RAG 검색 → 영상 재생 연결 (75줄)", "완료"],
            ["frontend/js/video_player.js", "Supabase URL 수신 → <video> 재생 플레이어", "완료"],
            ["frontend/js/vision.js", "카메라 MediaPipe 수어 인식", "완료"],
            ["frontend/js/education.js", "수어 학습 모드", "완료"],
        ],
        col_widths=[2.4, 2.8, 0.8],
        row_colors=[C_DONE_BG]*5
    )

    add_heading(doc, "3-3. 데이터 현황", 2)
    add_table(doc,
        ["소스", "수량", "상태"],
        [
            ["Supabase Storage sign_videos 버킷", "1,111개 MP4", "팀원 업로드 완료"],
            ["api/data/sign_video_urls.json", "3,695개 단어 + sldict URL", "로컬 인덱스 완료"],
            ["Supabase DB sign_lemma", "1,134개", "구축 완료"],
            ["Supabase DB sign_motion", "1,197개", "구축 완료"],
            ["Supabase DB sign_source_video", "1,111개", "구축 완료"],
        ],
        col_widths=[2.8, 1.4, 1.8],
        row_colors=[C_DONE_BG]*5
    )

    # ══════════════════════════════════════════
    # 3. 시스템 아키텍처
    # ══════════════════════════════════════════
    add_heading(doc, "4. 시스템 아키텍처", 1)

    add_heading(doc, "4-1. 전체 흐름", 2)
    add_code(doc,
        "[사용자]\n"
        "  ├─ 텍스트 입력 / 음성(STT)\n"
        "  │     → translator.js\n"
        "  │         → POST /api/sign-language/search  (RAG: FAISS 검색)\n"
        "  │             → keyword 반환\n"
        "  │                 → GET /api/video/url/{keyword}\n"
        "  │                     → supabase_video.py\n"
        "  │                         ┌─ Supabase Storage에 있음 → 즉시 공개 URL 반환\n"
        "  │                         └─ 없음 → sign_video_urls.json 조회\n"
        "  │                                     → sldict 다운로드 → Supabase 업로드\n"
        "  │                                         → 공개 URL 반환\n"
        "  │                     → video_player.js → <video> 재생\n"
        "  │\n"
        "  └─ 카메라 (WebSocket)\n"
        "        → /api/ws/vision  (MediaPipe → KNN 분류)\n"
        "            → itda:rag:result 이벤트\n"
        "                → ITDAVideoPlayer.play(keyword)\n"
        "\n"
        "[서버 구성]\n"
        "  Backend  : FastAPI  — http://localhost:8000\n"
        "  Frontend : Python HTTP Server — http://localhost:3000\n"
        "\n"
        "[외부 의존성]\n"
        "  Supabase Storage (sign_videos)  : 영상 CDN\n"
        "  sldict.korean.go.kr             : 원본 영상 소스 (폴백)\n"
        "  sign_video_urls.json            : 로컬 URL 인덱스 (3,695개)"
    )

    add_heading(doc, "4-2. 영상 서빙 상세 흐름", 2)
    add_table(doc,
        ["단계", "동작", "담당"],
        [
            ["1", "keyword 수신 (RAG 검색 결과)", "sign_language.py"],
            ["2", "Supabase Storage HEAD 요청 — 파일 존재 확인", "supabase_video.py"],
            ["3-A", "[캐시 히트] 공개 URL 즉시 반환", "supabase_video.py"],
            ["3-B", "[미스] sign_video_urls.json에서 sldict URL 조회", "supabase_video.py"],
            ["4", "sldict에서 MP4 다운로드 (aiohttp, 30초 타임아웃)", "supabase_video.py"],
            ["5", "Supabase Storage POST 업로드 (service key 인증)", "supabase_video.py"],
            ["6", "공개 URL 반환 → 프론트 video_player.js", "video_proxy.py"],
            ["7", "<video> 태그 src 설정 후 재생", "video_player.js"],
        ],
        col_widths=[0.5, 3.5, 2.0]
    )

    # ══════════════════════════════════════════
    # 4. 기술 스택
    # ══════════════════════════════════════════
    add_heading(doc, "5. 기술 스택", 1)
    add_table(doc,
        ["분류", "기술", "버전/비고"],
        [
            ["백엔드 프레임워크", "FastAPI", "0.109.2"],
            ["비동기 HTTP", "aiohttp", "3.9.3 — Supabase REST API 직접 호출"],
            ["AI 검색", "sentence-transformers + FAISS", "RAG 기반 수어 키워드 검색"],
            ["컴퓨터 비전", "MediaPipe", "카메라 손 랜드마크 추출"],
            ["분류기", "KNN (자체 구현)", "수어 동작 분류"],
            ["설정 관리", "pydantic-settings", "2.1.0 — .env 로드"],
            ["스토리지", "Supabase Storage", "sign_videos 버킷 — MP4 CDN"],
            ["데이터베이스", "Supabase (PostgreSQL)", "sign_lemma / sign_motion / sign_source_video"],
            ["프론트엔드", "Vanilla JS (ES Module)", "Three.js 제거, 순수 JS"],
            ["영상 재생", "HTML5 <video>", "MP4 스트리밍"],
            ["STT", "Web Speech API / 자체 어댑터", "음성→텍스트 변환"],
            ["서버 런타임", "uvicorn", "0.27.1"],
        ],
        col_widths=[1.6, 2.4, 2.0]
    )

    # ══════════════════════════════════════════
    # 5. API 설계
    # ══════════════════════════════════════════
    add_heading(doc, "6. API 설계", 1)

    add_heading(doc, "6-1. 현재 API 목록", 2)
    add_table(doc,
        ["메서드", "경로", "설명", "상태"],
        [
            ["GET",  "/api/health",                  "서버 상태 확인",                    "완료"],
            ["POST", "/api/sign-language/search",    "텍스트 → 수어 키워드 RAG 검색",     "완료"],
            ["GET",  "/api/video/url/{keyword}",     "수어 영상 Supabase 공개 URL 반환",  "완료"],
            ["WS",   "/api/ws/vision",               "카메라 실시간 수어 인식",            "완료"],
            ["POST", "/api/collect/*",               "KNN 학습 데이터 수집",              "완료"],
        ],
        col_widths=[0.8, 2.4, 2.2, 0.8],
        row_colors=[C_DONE_BG]*5
    )

    add_heading(doc, "6-2. /api/video/url/{keyword} 상세", 2)
    doc.add_paragraph("요청:")
    add_code(doc, "GET /api/video/url/안녕하세요")
    doc.add_paragraph("성공 응답 (200):")
    add_code(doc,
        '{\n'
        '  "keyword": "안녕하세요",\n'
        '  "url": "https://xxxx.supabase.co/storage/v1/object/public/sign_videos/안녕하세요.mp4"\n'
        '}'
    )
    doc.add_paragraph("실패 응답 (404):")
    add_code(doc,
        '{\n'
        '  "detail": "\'안녕하세요\'에 해당하는 영상을 찾을 수 없습니다"\n'
        '}'
    )

    add_heading(doc, "6-3. 추가 예정 API (선택)", 2)
    add_table(doc,
        ["메서드", "경로", "설명", "우선순위"],
        [
            ["GET", "/api/video/list", "업로드된 영상 전체 목록 조회", "낮음"],
            ["POST", "/api/video/batch-upload", "여러 키워드 일괄 업로드", "중간"],
        ],
        col_widths=[0.8, 2.0, 2.6, 0.8],
        row_colors=[C_TODO_BG]*2
    )

    # ══════════════════════════════════════════
    # 6. 데이터 전략
    # ══════════════════════════════════════════
    add_heading(doc, "7. 데이터 전략", 1)

    add_heading(doc, "7-1. 영상 서빙 전략", 2)
    add_table(doc,
        ["전략", "설명", "적용 시점"],
        [
            ["캐시 우선 (Cache-First)",
             "Supabase에 이미 업로드된 1,111개 — HEAD 확인 후 즉시 URL 반환",
             "현재 구현됨"],
            ["온디맨드 업로드 (On-Demand)",
             "없는 경우 sldict에서 다운로드 후 Supabase 업로드 (3,695개 커버)",
             "현재 구현됨"],
            ["사전 배치 업로드 (Pre-Batch)",
             "sign_video_urls.json 전체 3,695개를 미리 Supabase에 업로드해 온디맨드 지연 제거",
             "중기 과제"],
        ],
        col_widths=[1.8, 3.0, 1.4]
    )

    add_heading(doc, "7-2. 파일명 규칙", 2)
    add_table(doc,
        ["항목", "규칙", "예시"],
        [
            ["Storage 버킷", "sign_videos", "—"],
            ["파일명 생성 규칙", "keyword.replace('/', '_').replace(' ', '_') + '.mp4'", "안녕하세요.mp4"],
            ["한글 파일명", "한글 그대로 사용 (URL 인코딩은 aiohttp가 처리)", "날씨.mp4"],
        ],
        col_widths=[1.6, 3.2, 1.2]
    )
    add_callout(doc, "⚠ 중요: 팀원이 업로드한 파일명 규칙이 위와 동일한지 반드시 확인 필요.\n"
                     "다른 규칙(숫자 ID, 영문 변환 등)을 사용했다면 supabase_video.py의 _storage_path() 함수를 수정해야 함.",
               "FFF3CD")

    add_heading(doc, "7-3. 시나리오 커버리지", 2)
    add_table(doc,
        ["항목", "값"],
        [
            ["시나리오 대화", "20쌍 (농인↔청인)"],
            ["시나리오 단어 총계", "80개"],
            ["현재 커버 가능", "73개 (91.2%)"],
            ["미커버 7개", "날씨, 맑다, 배려, 서툴다, 음식, 청인, 화면"],
        ],
        col_widths=[2.0, 4.0]
    )

    add_heading(doc, "7-4. 미커버 7개 단어 처리 방안", 2)
    add_table(doc,
        ["단어", "처리 방안", "우선순위"],
        [
            ["날씨", "sign_video_urls.json 내 '날씨' 유사어('맑음', '흐림') 매핑 또는 수동 업로드", "높음"],
            ["맑다", "'맑음' 유사어 매핑 검토", "높음"],
            ["배려", "직접 영상 촬영 또는 sldict 추가 검색", "중간"],
            ["서툴다", "'어색하다' 유사어 매핑 검토", "중간"],
            ["음식", "'식사', '밥' 유사어 매핑 검토", "높음"],
            ["청인", "수어사전에 없을 수 있음 — '듣는 사람' 대체어 검토", "낮음"],
            ["화면", "'스크린', '모니터' 유사어 매핑 검토", "낮음"],
        ],
        col_widths=[1.0, 3.5, 1.0]
    )

    # ══════════════════════════════════════════
    # 7. 단계별 개발 계획
    # ══════════════════════════════════════════
    add_heading(doc, "8. 단계별 개발 계획", 1)

    # Phase 1
    add_heading(doc, "Phase 1 — 연결 및 검증", 2)
    add_callout(doc, "목표: 팀원이 업로드한 영상이 실제로 재생되는 것을 end-to-end 확인", "E8F5E9")
    add_table(doc,
        ["#", "작업", "파일", "상태"],
        [
            ["1", ".env 파일 생성 — SUPABASE_URL, SUPABASE_SERVICE_KEY 입력",
             ".env", "예정"],
            ["2", "팀원에게 파일명 규칙 확인\n(예: 안녕하세요.mp4 vs word_001.mp4)",
             "supabase_video.py", "예정"],
            ["3", "파일명 규칙 불일치 시 _storage_path() 함수 수정",
             "api/services/supabase_video.py", "예정"],
            ["4", "백엔드 서버 재시작\n(작업 관리자 → Python 종료 → uvicorn 재기동)",
             "—", "예정"],
            ["5", "<video> 태그에 autoplay 속성 추가",
             "frontend/index.html", "예정"],
            ["6", "end-to-end 테스트: '안녕하세요' 입력 → 영상 재생 확인",
             "—", "예정"],
        ],
        col_widths=[0.3, 2.8, 2.1, 0.8],
        row_colors=[C_TODO_BG]*6
    )

    # Phase 2
    add_heading(doc, "Phase 2 — 핵심 기능 완성", 2)
    add_callout(doc, "목표: 80개 시나리오 단어 100% 커버 + UI 완성도 확보", "E8F5E9")
    add_table(doc,
        ["#", "작업", "파일", "상태"],
        [
            ["1", "영상 로딩 스피너 추가 (로딩 중 시각 피드백)",
             "frontend/js/video_player.js", "예정"],
            ["2", "영상 없음(404) 시 한국어 안내 메시지 표시",
             "frontend/js/video_player.js", "예정"],
            ["3", "미커버 7개 단어 유사어 매핑 추가",
             "api/services/supabase_video.py", "예정"],
            ["4", "카메라 인식 → 영상 재생 플로우 end-to-end 테스트",
             "frontend/js/vision.js", "예정"],
            ["5", "STT 음성 입력 → 영상 재생 플로우 테스트",
             "frontend/js/translator.js", "예정"],
            ["6", "시나리오 20쌍 전체 대화 시뮬레이션 테스트",
             "—", "예정"],
        ],
        col_widths=[0.3, 2.8, 2.1, 0.8],
        row_colors=[C_TODO_BG]*6
    )

    # Phase 3
    add_heading(doc, "Phase 3 — 안정화 및 최적화", 2)
    add_callout(doc, "목표: 성능 및 안정성 확보, 사전 배치 업로드 완료", "E8F5E9")
    add_table(doc,
        ["#", "작업", "파일", "상태"],
        [
            ["1", "사전 배치 업로드 스크립트 작성\n(sign_video_urls.json 3,695개 전체 Supabase 업로드)",
             "tools/batch_upload.py (신규)", "예정"],
            ["2", "인메모리 URL 캐싱 — 동일 키워드 반복 요청 최적화",
             "api/services/supabase_video.py", "예정"],
            ["3", "sldict 다운로드 재시도 로직 강화 (3회 retry, exponential backoff)",
             "api/services/supabase_video.py", "예정"],
            ["4", "모바일 자동재생 대응 (playsinline 속성, muted 확인)",
             "frontend/index.html", "예정"],
            ["5", "Supabase 미설정 시 graceful degradation (플레이스홀더 유지)",
             "frontend/js/video_player.js", "예정"],
        ],
        col_widths=[0.3, 2.8, 2.1, 0.8],
        row_colors=[C_TODO_BG]*5
    )

    # Phase 4
    add_heading(doc, "Phase 4 — 완성 및 시연 준비 (최종)", 2)
    add_callout(doc, "목표: 시연 가능한 완성 버전 확보", "E8F5E9")
    add_table(doc,
        ["#", "작업", "파일", "상태"],
        [
            ["1", "전체 시나리오 80단어 100% 커버 확인",
             "—", "예정"],
            ["2", "UI/UX 최종 점검 (애니메이션, 전환 효과)",
             "frontend/index.html", "예정"],
            ["3", "백엔드 응답 시간 측정 및 개선 (목표: 1초 이내)",
             "—", "예정"],
            ["4", "배포 환경 설정 (환경변수, CORS 정리)",
             "api/core/config.py", "예정"],
            ["5", "README 및 시연 가이드 작성",
             "README.md", "예정"],
        ],
        col_widths=[0.3, 2.8, 2.1, 0.8],
        row_colors=[C_TODO_BG]*5
    )

    # 일정 요약
    add_heading(doc, "8-5. 단계 요약", 2)
    add_table(doc,
        ["Phase", "핵심 목표", "핵심 산출물"],
        [
            ["Phase 1 — 연결·검증", "end-to-end 영상 재생 확인", "영상 재생 동작 확인"],
            ["Phase 2 — 기능 완성", "80단어 커버, UI 완성",      "80단어 커버, UI 완성도 확보"],
            ["Phase 3 — 안정화",   "성능·안정성 확보",           "배치 업로드, 캐시, 재시도 로직"],
            ["Phase 4 — 시연 준비","시연 가능 버전 완성",         "100% 커버, 시연 가이드"],
        ],
        col_widths=[2.0, 2.0, 2.0]
    )

    # ══════════════════════════════════════════
    # 8. 미결 이슈 및 리스크
    # ══════════════════════════════════════════
    add_heading(doc, "9. 미결 이슈 및 리스크", 1)
    add_table(doc,
        ["이슈", "영향도", "원인", "해결 방법"],
        [
            ["파일명 규칙 불일치",
             "🔴 높음",
             "팀원이 업로드한 파일명이 _storage_path() 생성 규칙과 다를 수 있음",
             "팀원에게 파일명 형식 확인 후 _storage_path() 수정"],
            [".env 미설정",
             "🔴 높음",
             "SUPABASE_URL/SUPABASE_SERVICE_KEY 없으면 모든 영상 요청 실패",
             ".env 파일 생성 후 Supabase 대시보드에서 키 복사"],
            ["백엔드 좀비 프로세스",
             "🟡 중간",
             "포트 8000 점유로 새 서버 기동 불가",
             "작업 관리자에서 Python 프로세스 종료 후 재시작"],
            ["sldict 외부 서버 불안정",
             "🟡 중간",
             "sldict.korean.go.kr 서버가 느리거나 차단될 수 있음",
             "사전 배치 업로드로 의존성 제거 (Phase 3)"],
            ["미커버 7개 단어",
             "🟢 낮음",
             "sign_video_urls.json에 해당 단어 영상 없음",
             "유사어 매핑 또는 직접 영상 추가 (Phase 2)"],
            ["<video> 자동재생 차단",
             "🟢 낮음",
             "브라우저 정책상 muted 없으면 autoplay 차단",
             "autoplay + muted + playsinline 속성 추가"],
        ],
        col_widths=[1.6, 0.8, 2.0, 1.8]
    )

    # ══════════════════════════════════════════
    # 9. 역할 분담
    # ══════════════════════════════════════════
    add_heading(doc, "10. 역할 분담", 1)
    add_table(doc,
        ["역할", "담당", "완료 항목", "예정 항목"],
        [
            ["Supabase 영상 업로드",
             "인기동",
             "sign_videos 버킷 MP4 1,111개 업로드 완료",
             "파일명 규칙 공유, 추가 업로드"],
            ["백엔드 API 개발",
             "고병진",
             "video_proxy.py, supabase_video.py, 전체 라우터",
             "파일명 매핑 수정, 재시도 로직, 배치 업로드 스크립트"],
            ["프론트엔드 개발",
             "손영란",
             "video_player.js, translator.js, index.html 영상 영역",
             "autoplay 추가, 로딩 UI, 에러 메시지"],
            ["데이터 파이프라인",
             "손현정",
             "sign_video_urls.json, Supabase DB 구축",
             "미커버 7개 단어 처리"],
            ["테스트·QA",
             "전원 공동",
             "—",
             "시나리오 80단어 전체 테스트"],
        ],
        col_widths=[1.4, 0.9, 2.0, 1.9]
    )

    # ══════════════════════════════════════════
    # 10. 체크리스트
    # ══════════════════════════════════════════
    add_heading(doc, "11. 개발 완료 체크리스트", 1)

    add_heading(doc, "Phase 1 — 연결 검증", 2)
    items1 = [
        ("[ ]", ".env 파일 생성 (SUPABASE_URL, SUPABASE_SERVICE_KEY)"),
        ("[ ]", "팀원에게 Supabase 파일명 규칙 확인"),
        ("[ ]", "파일명 규칙 불일치 시 _storage_path() 수정"),
        ("[ ]", "백엔드 서버 재시작 (포트 8000)"),
        ("[ ]", "<video> 태그 autoplay 속성 추가"),
        ("[ ]", "'안녕하세요' 입력 → 영상 재생 end-to-end 확인"),
    ]
    for mark, text in items1:
        p = doc.add_paragraph(f"{mark}  {text}", style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, "Phase 2 — 기능 완성", 2)
    items2 = [
        ("[ ]", "영상 로딩 스피너 추가"),
        ("[ ]", "404 에러 시 한국어 안내 메시지"),
        ("[ ]", "미커버 7개 단어 유사어 매핑"),
        ("[ ]", "카메라 인식 → 영상 재생 플로우 테스트"),
        ("[ ]", "STT → 영상 재생 플로우 테스트"),
        ("[ ]", "시나리오 20쌍 전체 시뮬레이션"),
    ]
    for mark, text in items2:
        p = doc.add_paragraph(f"{mark}  {text}", style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, "Phase 3 — 안정화", 2)
    items3 = [
        ("[ ]", "사전 배치 업로드 스크립트 (3,695개 전체)"),
        ("[ ]", "인메모리 URL 캐싱"),
        ("[ ]", "sldict 재시도 로직 (3회 retry)"),
        ("[ ]", "모바일 자동재생 대응"),
        ("[ ]", "오프라인 폴백 처리"),
    ]
    for mark, text in items3:
        p = doc.add_paragraph(f"{mark}  {text}", style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, "Phase 4 — 시연 준비", 2)
    items4 = [
        ("[ ]", "80단어 100% 커버 확인"),
        ("[ ]", "UI/UX 최종 점검"),
        ("[ ]", "응답 시간 1초 이내 확인"),
        ("[ ]", "README 및 시연 가이드 작성"),
    ]
    for mark, text in items4:
        p = doc.add_paragraph(f"{mark}  {text}", style="List Bullet")
        p.runs[0].font.size = Pt(10)

    # 푸터
    doc.add_paragraph()
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run("ITDA (잇다) 팀  |  고병진  ·  인기동  ·  손영란  ·  손현정  |  VIDEO-ONLY MODE 기반").font.size = Pt(9)

    doc.save(OUTPUT)
    print(f"저장 완료: {OUTPUT}")


if __name__ == "__main__":
    build()
