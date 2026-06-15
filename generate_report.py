"""
ITDA 프로젝트 개발 보고서 생성기 (2026-05-21)
실행: python generate_report.py
출력: ITDA_개발보고서_20260521.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = "ITDA_개발보고서_20260521.docx"

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더 행
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "DCE6F1")
        cell._tc.get_or_add_tcPr().append(shading)

    # 데이터 행
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    # 열 너비
    if col_widths:
        for col_i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_i].width = Inches(width)

    doc.add_paragraph()
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    pPr.append(ind)
    doc.add_paragraph()


def build():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ITDA (잇다) 프로젝트")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("개발 보고서")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("보고 기준일: 2026-05-21\n농인↔청인 소통 보조 수어 번역 웹앱").font.size = Pt(11)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 1. 프로젝트 개요 ─────────────────────────────
    add_heading(doc, "1. 프로젝트 개요", 1)
    add_table(doc,
        ["항목", "내용"],
        [
            ["프로젝트명", "ITDA (잇다) — 농인↔청인 소통 보조 수어 번역 웹앱"],
            ["보고 기준일", "2026-05-21"],
            ["직전 보고서", "2026-05-20 (아바타_영상으로변환한내용.pdf)"],
            ["변경 계기", "3D 아바타 방식 폐기 → 실제 수어 영상 직접 재생 방식으로 전환"],
        ],
        col_widths=[1.5, 4.5]
    )

    # ── 2. 계획 변경 요약 ─────────────────────────────
    add_heading(doc, "2. 계획 변경 요약 — 이전 대비 현재", 1)

    add_heading(doc, "2-1. 이전 계획 (PDF 기준, 2026-05-20)", 2)
    doc.add_paragraph("핵심 설계: 아바타 + 영상 동시 재생 (DUAL MODE)")
    add_code_block(doc,
        "입력 텍스트\n"
        "  → RAG 검색 → 수어 키워드\n"
        "    → [좌] Three.js 3D 아바타 애니메이션 (joint 데이터 기반)\n"
        "    → [우] 원본 MP4 영상 (Supabase sign_videos 버킷)"
    )
    add_table(doc,
        ["항목", "현황"],
        [
            ["Supabase DB", "sign_lemma 1,134개 / sign_motion 1,197개 / sign_source_video 1,111개"],
            ["Storage 버킷", "sign_videos — MP4 1,111개 업로드 완료"],
            ["시나리오 커버", "80개 중 73개 (91.2%), 미커버 7개: 날씨, 맑다, 배려, 서툴다, 음식, 청인, 화면"],
            ["파이프라인", "extract.py → retarget.py → ingest.py → verify.py → batch.py"],
            ["추후 계획", "API가 joint 데이터 + video_url 동시 반환, 프론트에서 동시 재생"],
        ],
        col_widths=[2.0, 4.0]
    )

    add_heading(doc, "2-2. 현재 계획 (2026-05-21)", 2)
    doc.add_paragraph("핵심 설계: 영상 단독 재생 (VIDEO-ONLY MODE)")
    add_code_block(doc,
        "입력 텍스트\n"
        "  → RAG 검색 → 수어 키워드\n"
        "    → GET /api/video/url/{keyword}\n"
        "      → Supabase Storage 확인 (있으면 즉시 URL 반환)\n"
        "      → 없으면: sldict 다운로드 → Supabase 업로드 → URL 반환\n"
        "          → <video> 태그 재생"
    )

    # ── 3. 아키텍처 변경사항 ──────────────────────────
    add_heading(doc, "3. 아키텍처 변경사항", 1)

    add_heading(doc, "3-1. 제거된 구성요소", 2)
    add_table(doc,
        ["구성요소", "파일", "변경 내용"],
        [
            ["Three.js 렌더러", "frontend/js/avatar.js", "index.html 로드 제거"],
            ["모션 로더", "frontend/js/motion_loader_v3.js", "로드 제거"],
            ["NPY 로더", "frontend/js/motion_loader_npy.js", "로드 제거"],
            ["리타게팅", "frontend/js/retargeting.js", "로드 제거"],
            ["참조 영상 PiP", "frontend/js/ref_video.js", "로드 제거"],
            ["Three.js importmap", "frontend/index.html", "제거"],
            ["#three-canvas", "frontend/index.html", "제거"],
            ["아바타/스켈레톤 버튼", "frontend/index.html", "제거"],
            ["translator.js 아바타 코드", "—", "1,147줄 → ~75줄 (93% 감소)"],
        ],
        col_widths=[1.8, 2.4, 1.8]
    )

    add_heading(doc, "3-2. 신규 구성요소", 2)
    add_table(doc,
        ["구성요소", "파일", "역할"],
        [
            ["영상 플레이어", "frontend/js/video_player.js", "Supabase URL 수신 → <video> 재생"],
            ["영상 URL 라우터", "api/routers/video_proxy.py", "GET /api/video/url/{keyword}"],
            ["Supabase 연동 서비스", "api/services/supabase_video.py", "순수 aiohttp REST API (SDK 미사용)"],
            [".env.example", ".env.example", "Supabase 자격증명 템플릿"],
        ],
        col_widths=[1.8, 2.4, 1.8]
    )

    add_heading(doc, "3-3. 유지된 구성요소", 2)
    add_table(doc,
        ["구성요소", "파일", "역할"],
        [
            ["RAG 검색 엔진", "api/services/rag_engine.py", "FAISS + sentence-transformers"],
            ["수어 검색 라우터", "api/routers/sign_language.py", "/api/sign-language/search"],
            ["WebSocket 비전", "api/routers/ws_vision.py", "실시간 카메라 수어 인식"],
            ["STT 어댑터", "api/routers/stt.py", "음성 → 텍스트 변환"],
            ["데이터 수집", "api/routers/collect.py", "KNN 훈련 데이터 수집"],
            ["KNN 분류기", "api/services/knn_classifier.py", "카메라 수어 분류"],
        ],
        col_widths=[1.8, 2.4, 1.8]
    )

    # ── 4. API 설계 현황 ──────────────────────────────
    add_heading(doc, "4. API 설계 현황", 1)

    add_heading(doc, "4-1. 기존 API (유지)", 2)
    add_table(doc,
        ["메서드", "경로", "설명"],
        [
            ["GET", "/api/health", "서버 상태 확인"],
            ["POST", "/api/sign-language/search", "텍스트 → 수어 키워드 RAG 검색"],
            ["WS", "/api/ws/vision", "카메라 실시간 수어 인식"],
            ["POST", "/api/collect/*", "학습 데이터 수집"],
        ],
        col_widths=[1.0, 2.5, 2.5]
    )

    add_heading(doc, "4-2. 신규 API", 2)
    add_table(doc,
        ["메서드", "경로", "설명"],
        [["GET", "/api/video/url/{keyword}", "수어 영상 Supabase 공개 URL 반환"]],
        col_widths=[1.0, 2.5, 2.5]
    )
    doc.add_paragraph("응답 예시:")
    add_code_block(doc,
        '{\n'
        '  "keyword": "안녕하세요",\n'
        '  "url": "https://[project].supabase.co/storage/v1/object/public/sign_videos/안녕하세요.mp4"\n'
        '}'
    )
    doc.add_paragraph("처리 흐름:")
    steps = [
        "Supabase Storage HEAD 요청으로 파일 존재 확인",
        "존재: 즉시 공개 URL 반환 (캐시 히트)",
        "미존재: sign_video_urls.json 에서 sldict 원본 URL 조회 → 다운로드 → Supabase 업로드 → URL 반환",
        "조회 불가: 404 반환",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")
    doc.add_paragraph()

    # ── 5. 데이터 현황 ────────────────────────────────
    add_heading(doc, "5. 데이터 현황", 1)

    add_heading(doc, "5-1. 영상 데이터 소스", 2)
    add_table(doc,
        ["소스", "단어 수", "용도"],
        [
            ["api/data/sign_video_urls.json", "3,695개", "sldict 원본 URL 인덱스 (온디맨드 다운로드용)"],
            ["Supabase Storage sign_videos 버킷", "1,111개", "기존 세션에서 사전 업로드된 영상"],
        ],
        col_widths=[2.5, 1.0, 2.5]
    )

    add_heading(doc, "5-2. 버킷명 불일치 수정 (완료)", 2)
    add_table(doc,
        ["구분", "버킷명", "비고"],
        [
            ["이전 세션 (PDF 기준)", "sign_videos (언더스코어)", "기존 1,111개 업로드된 버킷"],
            ["현재 코드 (수정 후)", "sign_videos (언더스코어)", "supabase_video.py BUCKET 상수 수정 완료"],
        ],
        col_widths=[2.0, 2.0, 2.0]
    )

    add_heading(doc, "5-3. 시나리오 커버리지", 2)
    add_table(doc,
        ["항목", "값"],
        [
            ["시나리오 대화", "20쌍 (농인↔청인)"],
            ["시나리오 단어 총계", "80개"],
            ["커버 가능", "73개 (91.2%)"],
            ["미커버", "7개: 날씨, 맑다, 배려, 서툴다, 음식, 청인, 화면"],
        ],
        col_widths=[2.0, 4.0]
    )

    # ── 6. 미완료 항목 및 후속 조치 ──────────────────
    add_heading(doc, "6. 미완료 항목 및 후속 조치", 1)

    add_heading(doc, "6-1. 즉시 조치 필요 (블로커)", 2)
    add_table(doc,
        ["항목", "현황", "조치"],
        [
            ["Supabase 자격증명 설정", ".env 미설정", ".env 파일 생성 후 SUPABASE_URL, SUPABASE_SERVICE_KEY 입력"],
            ["버킷명 불일치 수정", "이번 세션 수정 완료", "supabase_video.py BUCKET = 'sign_videos'"],
            ["백엔드 서버 재시작", "좀비 프로세스 점유 (포트 8000)", "작업 관리자에서 Python 프로세스 종료 후 재시작"],
        ],
        col_widths=[1.8, 1.8, 2.4]
    )

    add_heading(doc, "6-2. 단기 과제", 2)
    add_table(doc,
        ["항목", "설명"],
        [
            ["미커버 7개 단어 처리", "sign_video_urls.json 내 유사어 매핑 또는 수동 영상 추가"],
            ["기존 데이터 연동 확인", "1,111개 기업로드 영상이 새 코드로 정상 서빙되는지 테스트"],
            ["영상 로딩 UI 개선", "로딩 스피너, 에러 메시지 한국어화"],
            ["오프라인 폴백", "Supabase 미설정 시 graceful degradation"],
        ],
        col_widths=[2.0, 4.0]
    )

    add_heading(doc, "6-3. 중기 과제", 2)
    add_table(doc,
        ["항목", "설명"],
        [
            ["사전 배치 업로드", "3,695개 전체 영상을 Supabase에 미리 업로드 (온디맨드 지연 제거)"],
            ["sldict 다운로드 안정성", "타임아웃/재시도 로직 강화"],
            ["캐시 레이어", "인메모리 URL 캐싱 (동일 키워드 반복 요청 최적화)"],
            ["모바일 대응", "자동재생 정책 대응 (muted + playsinline 속성 추가)"],
        ],
        col_widths=[2.0, 4.0]
    )

    # ── 7. 현재 시스템 구성도 ────────────────────────
    add_heading(doc, "7. 현재 시스템 구성도", 1)
    add_code_block(doc,
        "[사용자]\n"
        "  ├─ 텍스트 입력 / 음성(STT)\n"
        "  │     → translator.js\n"
        "  │         → POST /api/sign-language/search (RAG)\n"
        "  │             → GET /api/video/url/{keyword}\n"
        "  │                 ┌─ Supabase Storage에 있음 → 공개 URL 반환\n"
        "  │                 └─ 없음 → sldict 다운로드 → 업로드 → URL 반환\n"
        "  │                     → <video> 재생\n"
        "  │\n"
        "  └─ 카메라 (WebSocket)\n"
        "        → /api/ws/vision (MediaPipe → KNN)\n"
        "            → itda:rag:result 이벤트\n"
        "                → ITDAVideoPlayer.play(keyword)\n"
        "\n"
        "[외부 의존성]\n"
        "  Supabase Storage      : 영상 CDN 캐시\n"
        "  sldict.korean.go.kr   : 국립국어원 한국수어사전 원본 영상\n"
        "  sign_video_urls.json  : 로컬 URL 인덱스 (3,695개)"
    )

    # ── 8. 환경 설정 가이드 ───────────────────────────
    add_heading(doc, "8. 환경 설정 가이드", 1)
    add_code_block(doc,
        "# .env 파일 생성 (프로젝트 루트)\n"
        "SUPABASE_URL=https://[your-project-id].supabase.co\n"
        "SUPABASE_SERVICE_KEY=[service_role_key]\n\n"
        "# 백엔드 재시작 (포트 8000)\n"
        "python -m uvicorn api.main:app --reload --port 8000\n\n"
        "# 프론트엔드 서버 (포트 3000)\n"
        "python -m http.server 3000 --directory frontend"
    )

    # 푸터
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run(f"보고 기준: 2026-05-21  |  ITDA 개발팀").font.size = Pt(9)

    doc.save(OUTPUT)
    print(f"저장 완료: {OUTPUT}")


if __name__ == "__main__":
    build()
